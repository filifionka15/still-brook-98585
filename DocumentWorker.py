import PyPDF2
from PyPDF2.generic import DictionaryObject
import docx
import os
import re
import docx2pdf
import telebot
import pathlib


class DocumentWorker:  
    #Unique function thaht might use for parse any file with ${word/letters}
    def replaceSpecWord(self, file_path : str, replacementOn : DictionaryObject, user : telebot.types.User):

        head, tail = os.path.split(file_path)
        filename, file_extension = os.path.splitext(tail)
        doc = docx.Document(file_path)
        

        for variable_key, variable_value in replacementOn.items():
            for paragraph in doc.paragraphs:
                if variable_key in paragraph.text:
                    paragraph.text = paragraph.text.replace(variable_key, variable_value)

            for table in doc.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            if variable_key in paragraph.text:
                                paragraph.text = paragraph.text.replace(variable_key, variable_value)

        newFileName = filename + "_generated_for_" + user.username + ".docx"
        generated_path =  os.path.join(head, newFileName)
        doc.save(generated_path)
        return generated_path

    #Code below to the end might use for work with simple string (read from file and parse that information)
    def __readAndParse(self, file_path : str): #parse pdf or docx file and get text from it

        filename, file_extension = os.path.splitext(file_path)

        if file_extension == '.pdf': #PDF mode doesn't needed

            pdfFileObj = open(file_path, 'rb')
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

            text = ''

            for x in pdfReader.pages:
                text += x.extractText()

            return text

        elif file_extension == '.docx':

            doc = docx.Document(file_path)

            text = ''

            for x in doc.paragraphs:
                text += x.text;

            return text

    def __findAllSpecWord(self, text : str): #find all [X] words in text and return list of them
        result = []

        for match in re.finditer(r"\[[\w\s]+\]", text):
            result.append((match.span(), match.group()))

        return result


    def replaceCharacters(self, path : str,  replacementOn : list):

        text = self.__readAndParse(path)
        parsedText = self.__findAllSpecWord(text)

        counter = 0

        for x in parsedText:
            text = text.replace(x[1], replacementOn[counter][1])
            counter += 1
        
        return text

    def convertDocx2Pdf(path : str):
        head, tail = os.path.split(path)
        filename, file_extension = os.path.splitext(tail)
        newFileName = filename + '.pdf'
        generated_path =  os.path.join(head, newFileName)

        docx2pdf.convert(path, generated_path)

        return generated_path





